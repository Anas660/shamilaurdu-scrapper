import os
import glob
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def process_html_to_word(html_file_path):
    # Extract surah_id from filename
    filename = os.path.basename(html_file_path)
    surah_id = filename.split('_')[1]
    
    print(f"Processing Surah {surah_id}...")
    
    # Read HTML file
    with open(html_file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Get surah title
    title = soup.find('title').text.strip() if soup.find('title') else f"Surah {surah_id}"
    
    # Create a new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Add title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a separator
    separator = doc.add_paragraph()
    separator.add_run('_' * 80)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # SECTION 1: Arabic Text
    arabic_heading = doc.add_heading('Arabic Text', level=2)
    arabic_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Extract Arabic text
    ar_div = soup.find("div", class_="ar")
    if ar_div:
        # Get all spans excluding verse numbers (spans with class="nm")
        arabic_spans = []
        for span in ar_div.find_all("span"):
            if "nm" not in span.get("class", []):
                arabic_spans.append(span)
        
        # Add each Arabic verse with number
        for i, span in enumerate(arabic_spans):
            verse_number = i + 1
            
            # Add Arabic text paragraph with right-to-left alignment
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Arabic text
            arabic_run = p.add_run(span.get_text(strip=True))
            arabic_run.font.size = Pt(18)
            arabic_run.font.name = 'Arabic Typesetting'
            
            # Add verse number
            verse_num = p.add_run(f" ﴿{verse_number}﴾ ")
            verse_num.font.size = Pt(14)
            verse_num.font.name = 'Arabic Typesetting'
            
            # Set RTL direction - compatible way
            p._p.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    
    # Add a separator
    doc.add_paragraph().add_run('_' * 80).font.size = Pt(10)
    
    # SECTION 2: Urdu Translation
    translation_heading = doc.add_heading('Urdu Translation', level=2)
    translation_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Extract Urdu translations
    ur_div = soup.find("div", class_="ur")
    if ur_div:
        urdu_spans = ur_div.find_all("span")
        
        # Add each Urdu translation
        for i, span in enumerate(urdu_spans):
            verse_number = i + 1
            
            # Extract text (removing HTML tags but keeping the main text)
            urdu_text = ""
            for element in span.children:
                if element.name != 'a' and element.name != 'sup':
                    urdu_text += element.get_text(strip=True) if hasattr(element, 'get_text') else str(element)
            
            urdu_text = urdu_text.replace("\n", " ").strip()
            if not urdu_text:
                continue
                
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Add verse number with translation
            trans_run = p.add_run(f"{verse_number}- {urdu_text}")
            trans_run.font.size = Pt(14)
            trans_run.font.name = 'Jameel Noori Nastaleeq'
            
            # Set RTL direction - compatible way
            p._p.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    
    # Add a separator
    doc.add_paragraph().add_run('_' * 80).font.size = Pt(10)
    
    # SECTION 3: Tafseer Notes
    tafseer_heading = doc.add_heading('Tafseer Notes', level=2)
    tafseer_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Extract tafseer notes
    nt_div = soup.find("div", class_="nt")
    if nt_div:
        tafseer_paragraphs = nt_div.find_all("p")
        
        # Add each tafseer note
        for p_tag in tafseer_paragraphs:
            # Skip empty paragraphs
            if not p_tag.get_text(strip=True):
                continue
                
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Extract note number if available
            note_num = ""
            if p_tag.find('n'):
                note_num = p_tag.find('n').get_text(strip=True)
                tafseer_run = p.add_run(f"حاشیہ {note_num} ")
                tafseer_run.font.size = Pt(12)
                tafseer_run.font.name = 'Jameel Noori Nastaleeq'
                tafseer_run.bold = True
            
            # Extract and add the tafseer text
            tafseer_text = p_tag.get_text(strip=True)
            if note_num:
                # Remove the note number from the beginning of text
                tafseer_text = tafseer_text.replace(note_num, "", 1).strip()
                # Remove common separators like "-" or ":" after note number
                tafseer_text = tafseer_text.lstrip("- :")
            
            tafseer_content = p.add_run(tafseer_text)
            tafseer_content.font.size = Pt(12)
            tafseer_content.font.name = 'Jameel Noori Nastaleeq'
            
            # Set RTL direction - compatible way
            p._p.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    
    # Create output directory structure
    output_base_dir = os.path.join(os.path.dirname(html_file_path), "..", "word_surahs")
    if not os.path.exists(output_base_dir):
        os.makedirs(output_base_dir)
    
    # Create surah-specific directory
    surah_dir = os.path.join(output_base_dir, f"surah_{surah_id}")
    if not os.path.exists(surah_dir):
        os.makedirs(surah_dir)
    
    # Save the document
    output_file = os.path.join(surah_dir, f"surah_{surah_id}.docx")
    doc.save(output_file)
    print(f"   ↳ Document saved to {output_file}")
    
    return output_file

def process_all_html_files():
    # Path to the HTML files
    html_folder = os.path.join("html_files")
    
    # Ensure the path exists
    if not os.path.exists(html_folder):
        print(f"HTML folder '{html_folder}' not found.")
        return
    
    # Get all HTML files
    html_files = glob.glob(os.path.join(html_folder, "surah_*_html.txt"))
    
    if not html_files:
        print(f"No HTML files found in {html_folder}")
        return
    
    print(f"Found {len(html_files)} HTML files to process")
    
    # Sort files by surah number
    def get_surah_number(filepath):
        filename = os.path.basename(filepath)
        parts = filename.split('_')
        if len(parts) >= 2:
            try:
                return int(parts[1])
            except ValueError:
                return 0
        return 0
    
    html_files.sort(key=get_surah_number)
    
    # Process each file
    processed_files = []
    for html_file in html_files:
        try:
            output_file = process_html_to_word(html_file)
            processed_files.append(output_file)
        except Exception as e:
            print(f"Error processing {html_file}: {e}")
    
    print(f"\n✅ Processing complete! Created {len(processed_files)} Word documents")
    print("Documents saved in the 'word_surahs' directory")

if __name__ == "__main__":
    process_all_html_files()