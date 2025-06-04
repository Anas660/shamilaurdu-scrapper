import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import glob

def create_quran_word_document(json_file_path):
    # Load JSON data
    with open(json_file_path, 'r', encoding='utf-8') as file:
        surah_data = json.load(file)
    
    # Create a new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Add title
    title = doc.add_heading(f"{surah_data['surah_name']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a separator
    separator = doc.add_paragraph()
    separator.add_run('_' * 80)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # SECTION 1: Add all Arabic verses with numbers
    arabic_heading = doc.add_heading('Arabic Text', level=2)
    arabic_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for verse in surah_data['verses']:
        # Check if verse has content (some verses might be empty)
        if 'arabic' not in verse or not verse['arabic']:
            continue
            
        # Add Arabic text paragraph with right-to-left alignment
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Arabic text
        arabic_run = p.add_run(verse['arabic'])
        arabic_run.font.size = Pt(18)
        arabic_run.font.name = 'Arabic Typesetting'
        
        # Add verse number
        verse_num = p.add_run(f" ﴿{verse['verse_number']}﴾ ")
        verse_num.font.size = Pt(14)
        verse_num.font.name = 'Arabic Typesetting'
        
        # Set RTL direction - compatible way
        p._p.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    
    # Add a separator
    doc.add_paragraph().add_run('_' * 80).font.size = Pt(10)
    
    # SECTION 2: Add translations with reference numbers
    translation_heading = doc.add_heading('Urdu Translation', level=2)
    translation_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for verse in surah_data['verses']:
        # Check if verse has content
        if 'urdu' not in verse or not verse['urdu']:
            continue
            
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add verse number with translation
        trans_run = p.add_run(f"{verse['verse_number']}- {verse['urdu']}")
        trans_run.font.size = Pt(14)
        trans_run.font.name = 'Jameel Noori Nastaleeq'
        
        # Set RTL direction - compatible way
        p._p.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    
    # Add a separator
    doc.add_paragraph().add_run('_' * 80).font.size = Pt(10)
    
    # SECTION 3: Add tafseer by reference numbers
    tafseer_heading = doc.add_heading('Tafseer (Notes)', level=2)
    tafseer_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Create a dictionary to store unique tafseer notes
    unique_tafseer = {}
    
    for verse in surah_data['verses']:
        # Check if verse has tafseer content
        if 'tafseer' in verse and verse['tafseer'] and 'tafseer_refs' in verse and verse['tafseer_refs']:
            # Each verse may have multiple refs to the same tafseer
            for ref in verse['tafseer_refs']:
                # Only add unique tafseer entries
                if ref not in unique_tafseer:
                    unique_tafseer[ref] = verse['tafseer']
    
    # Now add all unique tafseer notes in order
    for ref in sorted(unique_tafseer.keys(), key=lambda x: int(x) if x.isdigit() else x):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add tafseer reference number and content
        tafseer_run = p.add_run(f"حاشیہ نمبر {ref}: ")
        tafseer_run.font.size = Pt(12)
        tafseer_run.font.name = 'Jameel Noori Nastaleeq'
        tafseer_run.bold = True
        
        # Add the actual tafseer text
        tafseer_text = p.add_run(unique_tafseer[ref])
        tafseer_text.font.size = Pt(12)
        tafseer_text.font.name = 'Jameel Noori Nastaleeq'
        
        # Set RTL direction - compatible way
        p._p.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    
    # Save the document
    output_dir = os.path.join(os.path.dirname(json_file_path), "..", "word_files")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    base_name = os.path.basename(json_file_path)
    base_name_without_ext = os.path.splitext(base_name)[0]
    output_file = os.path.join(output_dir, f"{base_name_without_ext}.docx")
    
    doc.save(output_file)
    print(f"Document saved as {output_file}")
    
    return output_file

if __name__ == "__main__":
    # Use raw string literal for the path to avoid backslash issues
    json_folder = r"d:\pixelpk projects\shamilaurdu-scrapper\jsons_ready"
    
    # Get all JSON files in the folder
    json_files = glob.glob(os.path.join(json_folder, "*.json"))
    
    if not json_files:
        print(f"No JSON files found in {json_folder}")
        exit(1)
    
    print(f"Found {len(json_files)} JSON files to process")
    
    # Sort files by surah number
    json_files.sort(key=lambda x: int(os.path.basename(x).split("_")[1].split(".")[0]))
    
    # Process each file
    processed_files = []
    for json_file in json_files:
        print(f"Processing {os.path.basename(json_file)}...")
        try:
            output_file = create_quran_word_document(json_file)
            processed_files.append(output_file)
        except Exception as e:
            print(f"Error processing {json_file}: {e}")
    
    print(f"\nProcessing complete! Created {len(processed_files)} Word documents")
    print("Documents saved in the 'word_files' directory")